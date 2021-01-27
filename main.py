from docx import Document
import pandas as pd
import os
import zipfile
import shutil


def word2pic(path, zip_path, tmp_path, store_path):
    """
    将文档中的图片进行提取
    :param path: 文档docx的路径
    :param zip_path: docx重命名为zip
    :param tmp_path: 中转图片文件夹
    :param store_path: 图片存储的路径
    :return:
    """
    if not os.path.exists(store_path):
        os.makedirs(store_path)
    # 将docx文件重命名为zip文件
    os.rename(path, zip_path)
    # 进行解压
    f = zipfile.ZipFile(zip_path, 'r')
    # 将图片提取并保存
    for file in f.namelist():
        f.extract(file, tmp_path)
    # 释放该zip文件
    f.close()

    # 将docx文件从zip还原为docx
    os.rename(zip_path, path)
    # 得到缓存文件夹中图片列表
    pic = os.listdir(os.path.join(tmp_path, 'word/media'))
    # 将图片复制到最终的文件夹中
    for i in pic:
        # 根据word的路径生成图片的名称
        new_name = path.replace('\\', '_')
        new_name = new_name.replace(':', '') + '_' + i
        shutil.copy(os.path.join(tmp_path + '/word/media', i), os.path.join(store_path, new_name))
    # 删除缓冲文件夹中的文件，用以存储下一次的文件
    for i in os.listdir(tmp_path):
        # 如果是文件夹则删除
        if os.path.isdir(os.path.join(tmp_path, i)):
            shutil.rmtree(os.path.join(tmp_path, i))



def tabel2csv(doc_path,save_path):
    """
    文档表格抽取文字成csv
    :param doc_path: 文档路径
    :param save_path: csv保存路径
    :return:
    """
    doc = Document(doc_path)
    tables = doc.tables  # 获取文件中的表格集

    all_imgs = []
    img_id = []
    det_name = []
    pip_id = []
    det_level = []
    num = 0
    for i, table in enumerate(tables):
        num += 1
        if table.rows[0].cells[0].text != '缺陷编号':
            continue
        all_imgs.append(table.rows[0].cells[2].paragraphs[0])
        for row in table.rows:
            row_txt = []
            for cell in row.cells:
                c = cell.text
                row_txt.append(c)
            if row_txt[0] == '缺陷编号':
                img_id.append(num - 2)
            elif row_txt[0] == '管道编号':
                pip_id.append(row_txt[1])
            elif row_txt[0] == '缺陷名称':
                det_name.append(row_txt[1])
            elif row_txt[0] == '严重程度（等级）':
                det_level.append(row_txt[1])
    df = pd.DataFrame({'缺陷编号': img_id, '管道编号': pip_id, '缺陷名称': det_name, '严重程度（等级）': det_level})
    df.to_csv(save_path, index=False)


if __name__ == '__main__':
    # 1.抽取表格信息
    doc_path = 'det.docx'
    save_path = 'data/detection.csv'
    tabel2csv(doc_path,save_path)

    # 2. 抽取文档中所有图片
    word2pic(doc_path, 'data/test.zip', 'data/tmp', 'data/pic')










